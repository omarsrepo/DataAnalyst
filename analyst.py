import streamlit as st
import pandas as pd
import requests
from requests.auth import HTTPBasicAuth
import json
from io import BytesIO
from docx import Document
import time

# Initialize session state for uploaded data and analysis results
st.session_state.setdefault("data", None)
st.session_state.setdefault("analysis_results", None)

# Sidebar for file upload
st.sidebar.title("Upload Dataset")
uploaded_file = st.sidebar.file_uploader("Choose a file", type=['xlsx', 'xls', 'csv'])

# Load data and display it in the app
if uploaded_file is not None:
    if uploaded_file.type != 'text/csv':
        data = pd.read_excel(uploaded_file)
    else:
        data = pd.read_csv(uploaded_file)
    st.session_state["data"] = data
    st.write("Dataset:")
    st.write(data)


# Function to add data for analysis to JamAI Base
def add_data_to_jamai(data, table_id):
    json_data = data.to_dict(orient='records')
    url = "https://app.jamaibase.com/api/v1/gen_tables/action/rows/add"
    payload = {
        "data": {"data": json_data},
        "stream": True,
        "table_id": table_id,
    }
    headers = {"accept": "application/json", "content-type": "application/json"}

    try:
        response = requests.post(
            url,
            json=payload,
            headers=headers,
            auth=HTTPBasicAuth("cisunit1@gmail.com", "Pp40907971989!"),  # Replace with actual credentials
            verify=False  # Remove `verify=False` for production
        )
        response.raise_for_status()
        st.success("Data added for analysis!")
    except requests.exceptions.RequestException as e:
        st.error(f"Error adding data: {e}")


# Function to retrieve analysis results from JamAI Base
def get_analysis_results(table_id):
    url = f"https://app.jamaibase.com/api/v1/gen_tables/action/{table_id}/rows?limit=100"
    headers = {"accept": "application/json"}

    try:
        response = requests.get(url, headers=headers, auth=HTTPBasicAuth("cisunit1@gmail.com", "Pp40907971989!"))
        response.raise_for_status()
        return response.json()
    except requests.exceptions.RequestException as e:
        st.error(f"Error retrieving results: {e}")
        return None


# Run analysis button
if st.button("Run Analysis"):
    if st.session_state.get("data") is not None:
        add_data_to_jamai(st.session_state["data"], 'DataAnalyst')

        # Retrieve the results from JamAI and store them
        results = get_analysis_results('DataAnalyst')
        if results:
            st.session_state["analysis_results"] = results
    else:
        st.error("Please upload a dataset first.")

# Display analysis results if available
if st.session_state.get("analysis_results"):
    analysis = st.session_state["analysis_results"]

    # Display textual insights
    st.subheader("Analysis Insights")
    st.write(analysis.get("textual_insights", "No insights found."))

    # Display visualizations
    st.subheader("Visualizations")
    for viz in analysis.get("visualizations", []):
        st.image(viz, use_column_width=True)  # Assuming `viz` is a URL or base64 image data


# Function to generate a downloadable Word report
def generate_report(analysis):
    doc = Document()
    doc.add_heading("Data Analysis Report", 0)

    # Add insights
    doc.add_heading("Insights", level=1)
    doc.add_paragraph(analysis.get("textual_insights", "No insights available."))

    # Add visualizations (if they are image URLs)
    for idx, viz in enumerate(analysis.get("visualizations", []), 1):
        doc.add_paragraph(f"Visualization {idx}")
        try:
            image_response = requests.get(viz)
            if image_response.status_code == 200:
                doc.add_picture(BytesIO(image_response.content))
        except requests.RequestException:
            doc.add_paragraph("Failed to load visualization.")

    # Save to a BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Download button for the analysis report
if st.session_state.get("analysis_results"):
    st.download_button(
        label="Download Report",
        data=generate_report(st.session_state["analysis_results"]),
        file_name="analysis_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
